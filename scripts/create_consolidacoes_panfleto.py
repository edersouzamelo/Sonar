from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "consolidacoes-panfleto"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "panfleto_modulo_consolidacoes_sonar.docx"

NAVY = "0B2545"
BLUE = "1769AA"
TEAL = "0E8F8F"
GREEN = "BCECCB"
DARK_GREEN = "1B7A3D"
AMBER = "F2B84B"
RED = "D94B4B"
INK = "1F2937"
MUTED = "64748B"
LIGHT = "F6F8FB"
BORDER = "D7DEE8"
WHITE = "FFFFFF"


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def hexrgb(value):
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=hexrgb(fill), outline=hexrgb(outline) if outline else None, width=width)


def text(draw, xy, value, size=24, fill=INK, bold=False, anchor=None):
    draw.text(xy, value, fill=hexrgb(fill), font=font(size, bold), anchor=anchor)


def pill(draw, xy, value, fill, fg=WHITE, w=None):
    x, y = xy
    f = font(20, True)
    bbox = draw.textbbox((0, 0), value, font=f)
    width = w or (bbox[2] - bbox[0] + 34)
    rounded(draw, (x, y, x + width, y + 36), 18, fill)
    draw.text((x + width / 2, y + 18), value, fill=hexrgb(fg), font=f, anchor="mm")


def screenshot_access():
    img = Image.new("RGB", (1600, 900), hexrgb("EEF3F8"))
    d = ImageDraw.Draw(img)
    rounded(d, (42, 42, 1558, 858), 28, WHITE, BORDER, 2)
    rounded(d, (42, 42, 350, 858), 28, NAVY)
    rounded(d, (72, 72, 320, 150), 18, "12385F")
    text(d, (100, 93), "SONAR", 32, WHITE, True)
    text(d, (100, 126), "Controle e inteligência", 18, "BFD6F0")

    menu = ["Painel", "Licitações", "Organizações Militares", "Classes", "Agenda", "Colosso"]
    y = 205
    for item in menu:
        selected = item == "Classes"
        rounded(d, (74, y, 318, y + 54), 14, "EAF6FF" if selected else NAVY, "EAF6FF" if selected else NAVY)
        text(d, (118, y + 15), item, 20, NAVY if selected else "D5E7FA", True if selected else False)
        y += 66

    text(d, (410, 92), "Classes de Suprimento", 36, NAVY, True)
    text(d, (410, 140), "Escolha a classe e acesse o módulo Consolidações pela área de trabalho.", 22, MUTED)

    classes = [
        ("Classe I", "Subsistência"),
        ("Classe II", "Material de intendência"),
        ("Classe III", "Combustíveis"),
        ("Classe IV", "Construção"),
        ("Classe V", "Armamento e munição"),
        ("Classe VI", "Engenharia"),
    ]
    x0, y0 = 410, 220
    for idx, (title, sub) in enumerate(classes):
        x = x0 + (idx % 3) * 360
        y = y0 + (idx // 3) * 185
        rounded(d, (x, y, x + 320, y + 145), 18, WHITE, BORDER, 2)
        text(d, (x + 26, y + 24), title, 24, NAVY, True)
        text(d, (x + 26, y + 58), sub, 18, MUTED)
        if title == "Classe II":
            rounded(d, (x + 24, y + 94, x + 270, y + 126), 16, TEAL)
            text(d, (x + 42, y + 100), "Abrir Consolidações", 18, WHITE, True)

    rounded(d, (410, 610, 1438, 785), 22, "E9F7F7", "B8E5E5", 2)
    text(d, (445, 642), "Caminho rápido", 28, NAVY, True)
    text(d, (445, 690), "Menu lateral > Classes > escolha a Classe > clique em Consolidações", 26, INK, True)
    text(d, (445, 735), "A mesma estrutura existe para as classes I a X, com dados isolados por classe.", 22, MUTED)
    path = OUT_DIR / "print_01_acesso.png"
    img.save(path)
    return path


def screenshot_grid():
    img = Image.new("RGB", (1800, 1050), hexrgb("F4F7FA"))
    d = ImageDraw.Draw(img)
    rounded(d, (35, 35, 1765, 1015), 22, WHITE, BORDER, 2)
    text(d, (80, 75), "Consolidações | Classe II", 38, NAVY, True)
    text(d, (80, 124), "Planilha operacional por OM, por grande comando, com anexos cumulativos.", 24, MUTED)
    pill(d, (1390, 75), "+ Nova coluna", TEAL, w=190)
    pill(d, (1598, 75), "Arquivadas", BLUE, w=130)

    x, y = 80, 190
    widths = [470, 270, 270, 270, 270]
    headers = ["Grande Cmdo / OM", "Tabela Tal\nPrazo: 30/06\nPor OM", "DIEx resposta\nPrazo: 05/07\nPor Cmdo", "Quantitativo\nPrazo: 10/07\nPor OM", "Nova demanda\nRenomeável"]
    row_h = 78
    for i, h in enumerate(headers):
        rounded(d, (x + sum(widths[:i]), y, x + sum(widths[: i + 1]), y + row_h), 0, "EAF0F7", BORDER, 2)
        for j, line in enumerate(h.split("\n")):
            text(d, (x + sum(widths[:i]) + 18, y + 12 + j * 20), line, 18 if j else 21, NAVY if j == 0 else MUTED, True if j == 0 else False)

    rows = [
        ("9ª Região Militar", "cmd", ["Recebido", "Recebido", "", ""]),
        ("Hospital Militar de Área de Campo Grande", "om", ["2 anexos", "Travada", "", ""]),
        ("Colégio Militar de Campo Grande", "om", ["Recebido", "Travada", "1 anexo", ""]),
        ("9º Grupamento Logístico", "cmd", ["Recebido", "Recebido", "Pendente", ""]),
        ("Companhia de Comando do 9º Gpt Log", "om", ["1 anexo", "Travada", "", ""]),
        ("18ª Brigada de Infantaria de Pantanal", "cmd", ["Pendente", "Pendente", "", ""]),
        ("9º CGCFEx", "om", ["Recebido", "Travada", "Recebido", ""]),
    ]
    y += row_h
    for r, (name, kind, values) in enumerate(rows):
        row_y = y + r * 86
        fill = "F8FAFC" if r % 2 == 0 else WHITE
        for c, w in enumerate(widths):
            cell_x = x + sum(widths[:c])
            rounded(d, (cell_x, row_y, cell_x + w, row_y + 86), 0, fill, BORDER, 1)
        if kind == "cmd":
            text(d, (x + 22, row_y + 25), "v  " + name, 23, NAVY, True)
        else:
            size = 17 if len(name) > 36 else 19
            text(d, (x + 46, row_y + 28), name, size, INK)
        for c, val in enumerate(values, start=1):
            cell_x = x + sum(widths[:c])
            if val == "Recebido" or "anexo" in val:
                rounded(d, (cell_x + 18, row_y + 20, cell_x + widths[c] - 18, row_y + 66), 16, GREEN, "90D9A6", 1)
                text(d, (cell_x + 36, row_y + 32), val, 19, DARK_GREEN, True)
            elif val == "Travada":
                rounded(d, (cell_x + 18, row_y + 20, cell_x + widths[c] - 18, row_y + 66), 16, "E5E7EB", "CDD3DA", 1)
                text(d, (cell_x + 42, row_y + 32), "Bloqueada por Cmdo", 18, MUTED, True)
            elif val == "Pendente":
                rounded(d, (cell_x + 18, row_y + 20, cell_x + widths[c] - 18, row_y + 66), 16, "FFF4D9", "F4CC75", 1)
                text(d, (cell_x + 42, row_y + 32), "Pendente", 18, "7A5A00", True)
            else:
                rounded(d, (cell_x + 72, row_y + 22, cell_x + widths[c] - 72, row_y + 64), 18, "EFF6FF", "BFDBFE", 1)
                text(d, (cell_x + widths[c] / 2, row_y + 43), "Upload", 18, BLUE, True, anchor="mm")

    rounded(d, (80, 890, 1720, 960), 18, "ECFDF5", "A7F3D0", 2)
    text(d, (110, 912), "Célula verde = arquivo recebido. Cada célula aceita anexos cumulativos, leitura/OCR e download.", 24, DARK_GREEN, True)
    path = OUT_DIR / "print_02_planilha.png"
    img.save(path)
    return path


def screenshot_colosso():
    img = Image.new("RGB", (1600, 900), hexrgb("F6F8FB"))
    d = ImageDraw.Draw(img)
    rounded(d, (42, 42, 1558, 858), 26, WHITE, BORDER, 2)
    rounded(d, (42, 42, 1558, 160), 26, NAVY)
    text(d, (88, 76), "Colosso", 42, WHITE, True)
    text(d, (88, 124), "Consulta inteligente das consolidações registradas no banco de dados", 22, "C7D7EA")

    rounded(d, (110, 220, 1120, 315), 24, "EAF2FF", "C9DDF7", 2)
    text(d, (145, 245), "Pergunta", 18, BLUE, True)
    text(d, (145, 275), "Quais OM ainda não enviaram a Tabela Tal da Classe II?", 27, NAVY, True)

    rounded(d, (330, 365, 1490, 690), 24, "F8FAFC", BORDER, 2)
    text(d, (370, 400), "Resposta do Colosso", 20, TEAL, True)
    lines = [
        "Com base nas consolidações da Classe II, a Tabela Tal já foi enviada por 5 OM.",
        "Ainda constam pendentes: 18ª Brigada de Infantaria de Pantanal e outras OM sem anexo na coluna.",
        "Posso filtrar por grande comando, por prazo, por tipo de documento ou por quantidade solicitada.",
    ]
    yy = 442
    for line in lines:
        text(d, (370, yy), line, 24, INK)
        yy += 52

    rounded(d, (370, 610, 1425, 655), 18, "E9F7F7", "B8E5E5", 1)
    text(d, (392, 622), "Exemplos: \"quanto a OM X pediu?\" | \"quem entregou o DIEx tal?\" | \"o que vence esta semana?\"", 20, NAVY, True)

    rounded(d, (110, 730, 1490, 805), 20, "FFF7E6", "F1D18A", 2)
    text(d, (145, 752), "Importante", 22, "7A5A00", True)
    text(d, (285, 752), "Arquivos enviados, OCR e metadados ficam disponíveis para consulta pelo RAG do SONAR.", 22, INK)
    path = OUT_DIR / "print_03_colosso.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D7DEE8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_para(paragraph, size=11, color=INK, bold=False, align=None, before=0, after=6):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_heading(doc, title, level=1):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18 if level == 1 else 14)
    run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, value, size=10.5, color=INK, bold=False):
    p = doc.add_paragraph()
    p.add_run(value)
    style_para(p, size=size, color=color, bold=bold, after=6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        style_para(p, size=10.5, color=INK, after=3)


def add_card_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.2)
    hdr = table.rows[0].cells
    hdr[0].text = "Funcionalidade"
    hdr[1].text = "O que muda na rotina"
    for cell in hdr:
        set_cell_shading(cell, "EAF0F7")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            style_para(p, 10.5, NAVY, True, after=0)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                style_para(p, 10, INK, False, after=0)
    doc.add_paragraph()


def add_image(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.5))
    p = doc.add_paragraph()
    p.add_run(caption)
    style_para(p, 9, MUTED, False, WD_ALIGN_PARAGRAPH.CENTER, after=8)


def build_doc():
    screenshots = [screenshot_access(), screenshot_grid(), screenshot_colosso()]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(0.75))

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SONAR | Módulo Consolidações")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Organize demandas por OM, acompanhe prazos e consulte a situação pelo Colosso.")
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p.paragraph_format.space_after = Pt(12)

    logo = ROOT / "public" / "sonar-logo-transparent.png"
    if logo.exists():
        doc.add_picture(str(logo), width=Inches(1.05))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(
        doc,
        "O módulo Consolidações transforma a coleta de informações das Classes de Suprimento em uma planilha viva: cada OM aparece na relação oficial, cada coluna representa uma demanda consolidável e cada célula guarda os arquivos enviados, com status visual e consulta pelo agente Colosso.",
        11.2,
        INK,
    )

    add_image(doc, screenshots[0], "Print demonstrativo: acesso pelo menu Classes e cartão Consolidações.")

    add_heading(doc, "Como acessar", 1)
    add_bullets(
        doc,
        [
            "Abra o SONAR e use o menu lateral Classes.",
            "Escolha a Classe de Suprimento desejada, de I a X.",
            "Clique em Consolidações na área de trabalho da classe.",
            "A estrutura é compartilhada entre as classes, mas os arquivos e dados ficam isolados por classe.",
        ],
    )

    add_heading(doc, "O que foi implementado", 1)
    add_card_table(
        doc,
        [
            ("Relação oficial de OM", "As OM vêm do módulo Organizações Militares. Ao editar a lista oficial, o restante do sistema espelha a mesma relação."),
            ("Grandes comandos recolhíveis", "As OM ficam agrupadas sob seus respectivos grandes comandos, com abertura e fechamento suave para facilitar a leitura."),
            ("Colunas renomeáveis", "Cada coluna pode representar uma nova demanda, documento, tabela ou informação a consolidar."),
            ("Prazo e escopo", "Toda coluna pode ter prazo e modo Por OM ou Por Cmdo. No modo Por Cmdo, apenas grandes comandos recebem upload."),
            ("Upload cumulativo", "Cada célula pode receber mais de um arquivo, com download posterior dos anexos enviados."),
            ("Status visual", "Ao receber arquivo, a célula fica verde clara e indica entrega com check ou contagem de anexos."),
            ("Arquivar ou eliminar", "Colunas podem ser arquivadas para sair da visualização ou eliminadas definitivamente com confirmação."),
            ("Agenda integrada", "Prazos cadastrados nas consolidações também aparecem na Agenda do SONAR."),
        ],
    )

    add_image(doc, screenshots[1], "Print demonstrativo: planilha por OM, prazos, escopo Por OM/Por Cmdo e células com anexos.")

    add_heading(doc, "Como trabalhar no dia a dia", 1)
    add_bullets(
        doc,
        [
            "Crie uma coluna para cada demanda que precisa ser consolidada.",
            "Informe o prazo para que a demanda também entre na Agenda.",
            "Defina se a entrega será Por OM ou Por Cmdo.",
            "Faça upload na célula correspondente à OM ou ao grande comando.",
            "Use arquivar para retirar demandas concluídas da visão principal sem apagar histórico.",
        ],
    )

    add_heading(doc, "Colosso entende as consolidações", 1)
    add_body(
        doc,
        "Os dados persistidos no banco, os metadados dos uploads e o texto extraído dos anexos ficam disponíveis para o RAG do SONAR. Assim, o Colosso pode responder perguntas operacionais sobre a situação das demandas, cruzando classe, OM, prazo, coluna e documentos enviados.",
        11,
    )
    add_bullets(
        doc,
        [
            "Quais OM ainda não enviaram a Tabela Tal?",
            "Quais unidades já entregaram o documento solicitado?",
            "Quanto de determinado item a OM X pediu, quando essa informação constar no arquivo ou na consolidação?",
            "Quais demandas vencem nesta semana?",
        ],
    )

    add_image(doc, screenshots[2], "Print demonstrativo: consulta do Colosso sobre situação das consolidações.")

    add_heading(doc, "Mensagem para o setor", 1)
    add_body(
        doc,
        "Use o módulo Consolidações como ponto único para registrar demandas, arquivos recebidos e pendências das Classes de Suprimento. Quanto mais organizada estiver a planilha, melhor o SONAR acompanha prazos e melhor o Colosso responde às perguntas do setor.",
        11.2,
        NAVY,
        True,
    )

    footer = section.footer.paragraphs[0]
    footer.text = "SONAR - Módulo Consolidações"
    style_para(footer, 9, MUTED, False, WD_ALIGN_PARAGRAPH.CENTER, after=0)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
